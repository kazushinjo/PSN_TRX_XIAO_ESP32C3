from docx import Document

# docx ファイルを開く
doc = Document("/Users/kazuichishinjo/Desktop/PSN_TRX_Nano_R4/docs/PSN_TRX_VFO改造仕様書.docx")

# 置換のペアリスト
replacements = [
    ("ESP32", "Arduino Nano R4"),
    ("GPIO21", "A4"),
    ("GPIO22", "A5"),
    ("GPIO32", "D2"),
    ("GPIO33", "D3"),
    ("GPIO25", "D4"),
    ("Preferences", "EEPROM"),
    ("NVS", "EEPROM"),
    ("IRAM_ATTR", ""),
    ("SDA=21, SCL=22", "A4=SDA, A5=SCL"),
    ("SDA  -> GPIO21", "SDA  -> A4"),
    ("SCL  -> GPIO22", "SCL  -> A5"),
    ("Encoder A   -> GPIO32", "Encoder A   -> D2"),
    ("Encoder B   -> GPIO33", "Encoder B   -> D3"),
    ("Encoder SW  -> GPIO25", "Encoder SW  -> D4"),
]

# 全段落と実行内容の置換処理
for paragraph in doc.paragraphs:
    for old, new in replacements:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

# テーブル内のテキストも置換
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for old, new in replacements:
                    if old in paragraph.text:
                        for run in paragraph.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)

# 保存
doc.save("/Users/kazuichishinjo/Desktop/PSN_TRX_Nano_R4/docs/PSN_TRX_VFO改造仕様書.docx")
print("docx file updated successfully!")
